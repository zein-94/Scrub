# Scrub CLI — Privacy-first PII masking for AI agents
# Copyright (C) 2024 Scrub Contributors
# Licensed under GNU GPL v3 — see LICENSE for details
# modes/document.py

from pathlib import Path
from typing import Optional

from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn

from vault import Vault
from detectors.pii import detect_and_mask


# ─────────────────────────────────────────────
# Supported file types
# ─────────────────────────────────────────────

SUPPORTED_TYPES = {
    "txt":  "Plain Text",
    "md":   "Markdown",
    "docx": "Word Document",
    "xlsx": "Excel Spreadsheet",
    "csv":  "CSV File",
    "pdf":  "PDF Document",
    "eml":  "Email File",
    "html": "HTML File",
    "json": "JSON File",
    "xml":  "XML File",
}


# ─────────────────────────────────────────────
# File readers
# ─────────────────────────────────────────────

def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    except Exception as e:
        raise RuntimeError(f"Failed to read .docx: {e}")


def _read_xlsx(path: Path) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True)
        parts = []

        for sheet in wb.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(
                    str(cell) for cell in row if cell is not None and str(cell).strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    except Exception as e:
        raise RuntimeError(f"Failed to read .xlsx: {e}")


def _read_csv(path: Path) -> str:
    import csv
    rows = []
    with path.open(encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(cell.strip() for cell in row if cell.strip()))
    return "\n".join(rows)


def _read_pdf(path: Path) -> str:
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
        return "\n".join(parts)

    except Exception as e:
        raise RuntimeError(f"Failed to read .pdf: {e}")


def _read_eml(path: Path) -> str:
    import email
    from email import policy

    raw = path.read_bytes()
    msg = email.message_from_bytes(raw, policy=policy.default)
    parts = []

    # Headers
    for header in ("From", "To", "Cc", "Bcc", "Subject", "Date", "Reply-To"):
        val = msg.get(header)
        if val:
            parts.append(f"{header}: {val}")

    parts.append("")  # blank line between headers and body

    # Body — prefer plain text, fall back to html
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == "text/plain":
                payload = part.get_payload(decode=True)
                if payload:
                    parts.append(payload.decode("utf-8", errors="replace"))
                    break
        else:
            # Fallback to HTML if no plain text part found
            for part in msg.walk():
                if part.get_content_type() == "text/html":
                    payload = part.get_payload(decode=True)
                    if payload:
                        import re
                        html = payload.decode("utf-8", errors="replace")
                        parts.append(re.sub(r"<[^>]+>", " ", html))
                    break
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            parts.append(payload.decode("utf-8", errors="replace"))

    return "\n".join(parts)


def _read_html(path: Path) -> str:
    import re
    html = path.read_text(encoding="utf-8", errors="replace")
    # Strip tags, collapse whitespace
    text = re.sub(r"<[^>]+>", " ", html)
    text = re.sub(r"\s{2,}", "\n", text)
    return text.strip()


def _read_json(path: Path) -> str:
    import json
    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    return json.dumps(data, indent=2, ensure_ascii=False)


def _read_xml(path: Path) -> str:
    import xml.etree.ElementTree as ET

    tree = ET.parse(str(path))
    root = tree.getroot()
    parts = []

    def _walk(node: ET.Element, depth: int = 0):
        indent = "  " * depth
        text = (node.text or "").strip()
        attribs = " ".join(f'{k}="{v}"' for k, v in node.attrib.items())
        header = f"{indent}<{node.tag}{' ' + attribs if attribs else ''}>"
        if text:
            parts.append(f"{header} {text}")
        else:
            parts.append(header)
        for child in node:
            _walk(child, depth + 1)

    _walk(root)
    return "\n".join(parts)


# ─────────────────────────────────────────────
# Reader dispatch
# ─────────────────────────────────────────────

READERS = {
    "txt":  _read_txt,
    "md":   _read_txt,
    "docx": _read_docx,
    "xlsx": _read_xlsx,
    "csv":  _read_csv,
    "pdf":  _read_pdf,
    "eml":  _read_eml,
    "html": _read_html,
    "json": _read_json,
    "xml":  _read_xml,
}


def _read_file(path: Path, file_type: str, console: Console) -> str:
    """Dispatch to the correct reader based on file extension."""
    reader = READERS.get(file_type)

    if reader is None:
        console.print(
            f"[yellow]⚠[/yellow]  Unsupported file type [bold].{file_type}[/bold]. "
            f"Treating as plain text."
        )
        return _read_txt(path)

    return reader(path)


# ─────────────────────────────────────────────
# Document mode entry point
# ─────────────────────────────────────────────

def run_document_mode(
    input_path: Optional[Path],
    raw_text: Optional[str],
    file_type: str,
    language: str,
    console: Console,
) -> tuple[str, Vault]:
    """
    Full document mode pipeline:
      1. Read file (if not already raw text)
      2. Detect + mask PII
      3. Return masked text + vault
    """
    vault = Vault()

    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        console=console,
        transient=True,
    ) as progress:

        # ── Step 1: Read ─────────────────────────────────────────────
        task = progress.add_task("Reading file...", total=None)

        if raw_text is not None:
            text = raw_text
            file_label = "stdin"
        else:
            file_label = SUPPORTED_TYPES.get(file_type, f".{file_type} file")
            console.print(
                f"\n[cyan]→[/cyan] Reading [bold]{file_label}[/bold]: "
                f"[dim]{input_path}[/dim]"
            )
            try:
                text = _read_file(input_path, file_type, console)
            except RuntimeError as e:
                console.print(f"[red]✗[/red] {e}")
                raise SystemExit(1)

        progress.update(task, description="File loaded.")

        if not text.strip():
            console.print("[yellow]⚠[/yellow]  File appears to be empty.")
            return "", vault

        word_count  = len(text.split())
        char_count  = len(text)
        console.print(
            f"[dim]  {word_count:,} words · {char_count:,} characters[/dim]"
        )

        # ── Step 2: PII Detection + Masking ──────────────────────────
        progress.update(task, description="Scanning for PII...")
        console.print(
            f"\n[cyan]→[/cyan] Running PII detection "
            f"[dim](language: {language})[/dim]"
        )

        masked_text = detect_and_mask(
            text=text,
            vault=vault,
            language=language,
        )

        progress.update(task, description="Done.")

    # ── Results summary ───────────────────────────────────────────────
    if not vault:
        console.print("\n[green]✓[/green] No PII detected in document.")
    else:
        # Group by type for a clean breakdown
        type_counts: dict[str, int] = {}
        for _, entry in vault.items():
            t = entry["type"]
            type_counts[t] = type_counts.get(t, 0) + 1

        console.print(
            f"\n[green]✓[/green] Detected [bold]{len(vault)}[/bold] "
            f"sensitive item(s) across "
            f"[bold]{len(type_counts)}[/bold] type(s):"
        )
        for entity_type, count in sorted(type_counts.items()):
            console.print(f"  [yellow]·[/yellow] {entity_type}: [bold]{count}[/bold]")

    return masked_text, vault